import gensim
import nltk
import PyPDF2
import textract
import docx
import numpy as np
from nltk.tokenize import word_tokenize, sent_tokenize

def similarity(file_to_be_checked, file_to_check_against):
    file_docs = []
    file2_docs = []
    avg_sims = []

    doc = docx.Document('core/media/' + file_to_be_checked)  # Creating word reader object.
    data = ""
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)

    # print(data)
    # document = docx2txt.process('core/media/' + file_to_be_checked)
    # with open('core/media/documents/' + file_to_check_against) as f:
    # with open(document) as f:
    tokens = sent_tokenize(data)
    for line in tokens:
        file_docs.append(line)

    length_doc1 = len(file_docs)

    gen_docs = [[w.lower() for w in word_tokenize(text)]
                for text in file_docs]

    dictionary = gensim.corpora.Dictionary(gen_docs)
    corpus = [dictionary.doc2bow(gen_doc) for gen_doc in gen_docs]
    tf_idf = gensim.models.TfidfModel(corpus)
    sims = gensim.similarities.Similarity('core/workdir/', tf_idf[corpus],
                                          num_features=len(dictionary))

    doc = docx.Document('core/media/documents/' + file_to_check_against)  # Creating word reader object.
    data2 = ""
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        data2 = '\n'.join(fullText)
    # document2 = docx2txt.process('core/media/documents/' + file_to_check_against)
    #with open('core/media/' + file_to_be_checked) as f:
    #with open(document2) as f:
    tokens2 = sent_tokenize(data2)
    for line in tokens2:
        file2_docs.append(line)

    for line in file2_docs:
        query_doc = [w.lower() for w in word_tokenize(line)]
        query_doc_bow = dictionary.doc2bow(query_doc)
        query_doc_tf_idf = tf_idf[query_doc_bow]
        print('Comparing Result:', sims[query_doc_tf_idf])
        sum_of_sims = (np.sum(sims[query_doc_tf_idf], dtype=np.float32))
        avg = sum_of_sims / len(file_docs)
        print(f'avg: {sum_of_sims / len(file_docs)}')
        avg_sims.append(avg)
    total_avg = np.sum(avg_sims, dtype=np.float)
    print(total_avg)
    percentage_of_similarity = round(float(total_avg) * 100)
    if percentage_of_similarity >= 100:
        percentage_of_similarity = 100

    return percentage_of_similarity
